import os
import glob


def _find_mapping_file(mapping_folder, basename):
    expected_filename = f"ref_mapping_{basename}.txt"
    for filename in os.listdir(mapping_folder):
        if filename.lower() == expected_filename.lower():
            return os.path.join(mapping_folder, filename)
    return None


def process_files(file_folder, mapping_folder, remove_header=False):
    try:
        if not os.path.exists(file_folder):
            return "error", f"Files folder does not exist: {file_folder}"
        if not os.path.exists(mapping_folder):
            return "error", f"Mapping ref folder does not exist: {mapping_folder}"

        docx_files = glob.glob(os.path.join(file_folder, "*.docx"))
        doc_files = glob.glob(os.path.join(file_folder, "*.doc"))
        word_files = docx_files + doc_files

        if not word_files:
            return "no_file", "No file is found."

        missing_mapping_files = []
        processed_count = 0

        for word_file in word_files:
            basename = os.path.splitext(os.path.basename(word_file))[0]
            mapping_file = _find_mapping_file(mapping_folder, basename)

            if not mapping_file:
                missing_mapping_files.append(basename)
                continue

            mappings = read_mapping_file(mapping_file)
            if mappings:
                process_word_file(word_file, mappings, remove_header)
                processed_count += 1

        if missing_mapping_files:
            msg_lines = [f"No ref_mapping_{name}.txt is found." for name in missing_mapping_files]
            return "missing_mapping", "\n".join(msg_lines)

        return "success", "Completed."

    except Exception as e:
        return "error", str(e)


def read_mapping_file(mapping_file):
    mappings = []
    with open(mapping_file, 'r', encoding='utf-8') as f:
        for line_num, line in enumerate(f, 1):
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            if '<<<' not in line:
                continue

            parts = line.split('<<<', 1)
            if len(parts) != 2:
                raise ValueError(f"Invalid mapping line {line_num}: {line}")

            sTo = parts[0].strip()
            sFrom = parts[1].strip()

            if not sTo or not sFrom:
                raise ValueError(f"Invalid mapping line {line_num}: {line}")

            mappings.append((sTo, sFrom))

    return mappings


def process_word_file(word_file, mappings, remove_header=False):
    ext = os.path.splitext(word_file)[1].lower()
    basename = os.path.splitext(os.path.basename(word_file))[0]
    output_file = os.path.join(os.path.dirname(word_file), f"{basename}-masked{ext}")

    if ext == '.docx':
        process_docx(word_file, output_file, mappings, remove_header)
    elif ext == '.doc':
        process_doc(word_file, output_file, mappings, remove_header)


def process_docx(input_path, output_path, mappings, remove_header=False):
    from docx import Document

    doc = Document(input_path)

    if remove_header:
        _remove_headers_docx(doc)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, mappings)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, mappings)

    doc.save(output_path)


def _remove_headers_docx(doc):
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            for paragraph in header.paragraphs:
                paragraph.text = ""


def _replace_in_paragraph(paragraph, mappings):
    if not paragraph.runs:
        return

    full_text = paragraph.text
    if not full_text:
        return

    for sTo, sFrom in mappings:
        if sFrom not in full_text:
            continue

        if sTo.find(sFrom) != -1:
            idx = full_text.rfind(sFrom)
            while idx >= 0:
                match_start = idx
                match_end = idx + len(sFrom)

                run_start_positions = []
                char_offset = 0
                for run in paragraph.runs:
                    run_start_positions.append(char_offset)
                    char_offset += len(run.text)

                first_run_idx = None
                last_run_idx = None
                for i, run in enumerate(paragraph.runs):
                    run_start = run_start_positions[i]
                    run_end = run_start + len(run.text)
                    if run_start <= match_start < run_end:
                        first_run_idx = i
                    if run_start < match_end <= run_end:
                        last_run_idx = i
                        break

                if first_run_idx is None or last_run_idx is None:
                    idx = full_text.rfind(sFrom, 0, match_start)
                    continue

                _merge_and_replace(paragraph.runs, first_run_idx, last_run_idx,
                                 match_start - run_start_positions[first_run_idx],
                                 match_end - run_start_positions[last_run_idx], sTo)
                full_text = paragraph.text

                idx = full_text.rfind(sFrom, 0, match_start)
        else:
            idx = full_text.find(sFrom)
            while idx != -1:
                match_start = idx
                match_end = idx + len(sFrom)

                run_start_positions = []
                char_offset = 0
                for run in paragraph.runs:
                    run_start_positions.append(char_offset)
                    char_offset += len(run.text)

                first_run_idx = None
                last_run_idx = None
                for i, run in enumerate(paragraph.runs):
                    run_start = run_start_positions[i]
                    run_end = run_start + len(run.text)
                    if run_start <= match_start < run_end:
                        first_run_idx = i
                    if run_start < match_end <= run_end:
                        last_run_idx = i
                        break

                if first_run_idx is None or last_run_idx is None:
                    break

                _merge_and_replace(paragraph.runs, first_run_idx, last_run_idx,
                                 match_start - run_start_positions[first_run_idx],
                                 match_end - run_start_positions[last_run_idx], sTo)

                full_text = paragraph.text
                idx = full_text.find(sFrom, match_end)


def _merge_and_replace(runs, first_idx, last_idx, start_in_first, end_in_last, sTo):
    before = ''.join(runs[i].text for i in range(first_idx)) + runs[first_idx].text[:start_in_first]

    if last_idx == first_idx:
        after = runs[last_idx].text[end_in_last:]
    else:
        after = runs[last_idx].text[end_in_last:] + ''.join(runs[i].text for i in range(last_idx + 1, len(runs)))

    runs[first_idx].text = before + sTo + after

    for i in range(first_idx + 1, last_idx + 1):
        runs[i].text = ""

    if first_idx > 0:
        for i in range(first_idx):
            runs[i].text = ""

    for i in range(first_idx + 1, last_idx + 1):
        runs[i].text = ""


def process_doc(input_path, output_path, mappings, remove_header=False):
    import win32com.client
    import pythoncom

    pythoncom.CoInitialize()

    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        doc = word.Documents.Open(os.path.abspath(input_path))

        if remove_header:
            _remove_headers_doc(doc)

        for sTo, sFrom in mappings:
            find_and_replace(word, sFrom, sTo)

        doc.SaveAs(os.path.abspath(output_path), FileFormat=16)
        doc.Close()
        word.Quit()

    finally:
        pythoncom.CoUninitialize()


def _remove_headers_doc(doc):
    from win32com.client import constants
    try:
        wdHeaderFooter = constants.wdHeaderFooter
    except:
        import win32com.client.constants
        wdHeaderFooter = win32com.client.constants.wdHeaderFooter

    for section in doc.Sections:
        for header_type in [wdHeaderFooter.wdHeaderFooterAllPages,
                           wdHeaderFooter.wdHeaderFooterFirst,
                           wdHeaderFooter.wdHeaderFooterPrimary]:
            try:
                header = section.Headers.Item(header_type)
                header.Range.Text = ""
            except:
                pass


def find_and_replace(word_app, sFrom, sTo):
    word_app.Selection.Find.ClearFormatting()
    word_app.Selection.Find.Text = sFrom
    word_app.Selection.Find.Replacement.Text = sTo
    word_app.Selection.Find.Forward = True
    word_app.Selection.Find.Wrap = 1
    word_app.Selection.Find.Format = False
    word_app.Selection.Find.MatchCase = False
    word_app.Selection.Find.MatchWholeWord = False

    word_app.Selection.Find.Execute(Replace=2)